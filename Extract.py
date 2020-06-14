##################################################################################
###  该代码的目的：读取word文件（含有文本块、表格或图片）和excel文件（含有文本块或图片）
###  第一步：读入外部文件
###  首先，会把含有图片的文件中的图片分离出来；然后分别读取文本块和表格。
###  1、word、excel中有隐藏的图片，excel中有隐藏的sheet表。
###  2、所有导入文件导出的结构化数据为：业务单位、文件类名称、文件名称、活动名称、品类、参与门店、活动时间、活动要求、活动支持、活动类型。
###                                                          编写人：HYR
###                                                         2018年5月28日
##################################################################################

# -*- coding: utf-8 -*-
import win32com.client as win32
import glob
import os
import zipfile
import pandas as pd
from pandas import Series, DataFrame
import xlrd
from docx import Document
import math


class readexcel(object):
    def __init__(self, path):
        self.path = path

    def read_marketingfile(self):
        marketingword = pd.DataFrame(columns=['参与门店', '活动时间', '活动要求', '活动支持'])
        word = []
        readmarketing = pd.read_excel(self.path, encoding='gbk', sheet_name='促销方案', header=None, sep='\t')
        for i in range(readmarketing.shape[0]):
            if str(readmarketing[0][i]).strip() == '执行要求' and str(readmarketing[1][i]).strip() == '门店类型':
                word.append(str(readmarketing[2][i]).strip())
                shoprow = i
            else:
                pass

        word.append(str(readmarketing[1][1]).strip())

        req_word = ''
        marketingfillnan = readmarketing.fillna(' ')
        for j in range(marketingfillnan.shape[0]):
            if str(marketingfillnan[0][j]) == ' ' and (str(marketingfillnan[1][j]).strip() == '堆码要求' or str(
                    marketingfillnan[1][j]).strip() == '门店布置要求'):
                req_word = req_word + str(marketingfillnan[2][j]).replace(' ', '') + '  '

        word.append(str(req_word).strip())

        sup_word = ''
        for k in range(3, shoprow - 1):
            sup_word = sup_word + str(readmarketing[2][k]).replace(' ', '') + '  '

        word.append(str(sup_word).strip())

        word.append(str(readmarketing[3][2]))

        marketingword.loc[0] = word

        return marketingword


    def re_xls_xlsx(self):
		file = os.path.split(self.path)[1]
		
		if not os.path.exists(self.path):
			print('No such File! :%s' % self.path)
			return False
		elif file.endswith('.xls'):
			excel = win32.gencache.EnsureDispatch('Excel.Application')
			wb = excel.Workbooks.Open(self.path)
			wb.SaveAs(self.path + "x", FileFormat=51)
			wb.Close()
			excel.Application.Quit()
		else:
			pass
			

    def getexceldata(self, selfdefine):
        file = os.path.split(self.path)[1]
        if file.endswith('.xlsx'):
            f_path = os.path.join(self.path, file)
            workbook = xlrd.open_workbook(f_path)
            sheetnames = workbook.sheet_names()
            print(sheetnames)

            sheet = workbook.sheet_by_name(selfdefine)
            listmergecell = sheet.merged_cells
            # print(listmergecell)
            readetf = pd.read_excel(f_path, encoding='gbk', sheet_name=selfdefine, header=1, sep='\t')
            # print(readetf)
            fname = str(os.path.splitext(file)[0]) + selfdefine + '.csv'
            fname_txt = str(os.path.splitext(file)[0]) + selfdefine + '.txt'
            print(readetf.columns)
            if listmergecell != []:
                for k in range(len(listmergecell)):
                    if listmergecell[k][0] < 2 and listmergecell[k][1] >= 2:
                        for i in range(2, listmergecell[k][1]):
                            for j in range(listmergecell[k][2], listmergecell[k][3]):
                                readetf.iloc[i - 2, j] = sheet.cell_value(listmergecell[k][0], listmergecell[k][2])
                    elif listmergecell[k][0] >= 2 and listmergecell[k][1] >= 2:
                        for i in range(listmergecell[k][0], listmergecell[k][1]):
                            for j in range(listmergecell[k][2], listmergecell[k][3]):
                                readetf.iloc[i - 2, j] = sheet.cell_value(listmergecell[k][0], listmergecell[k][2])
                    elif listmergecell[k][1] < 2:
                        print("###########提示：该合并的单元格为标题表头#########")
                data = readetf.drop_duplicates()
            #		data.to_csv(os.path.join(self.path, fname), encoding='utf_8_sig', index=False)
            #		data.to_csv(os.path.join(self.path, fname_txt), encoding='utf_8_sig', index=False)
            else:
                data = readetf.drop_duplicates()
                #		data.to_csv(os.path.join(self.path, fname), encoding='utf_8_sig', index=False)
                #		data.to_csv(os.path.join(self.path, fname_txt), encoding='utf_8_sig', index=False)
            data_rename = data.drop(columns=['报销条件', '费用来源'], axis=1).rename(
                    columns={'商店对象': '参与门店', '时间': '活动时间', '要求': '活动要求', '支持': '活动支持'})
 #       return data_rename


	def chage_file_name(self, old_type='.xlsx', new_type='.zip'):
		file = os.path.split(self.path)[1]
		
		if not os.path.exists(self.path):
			print('No such File! :%s' % self.path)
			return False
		elif file.endswith(old_type):
			new_name = str(os.path.splitext(file)[0]) + new_type
			new_path = os.path.join(self.path, new_name)
			if os.path.exists(new_path):
				os.remove(new_path)
			os.rename(self.path, new_path)
		else:
			pass
			

    def unzip_file(self):
        file = os.path.split(self.path)[1]
		if os.path.splitext(file)[1] == '.zip':
			file_zip = zipfile.ZipFile(os.path.join(self.path, file), 'r')
			zipdir = os.path.splitext(file)[0]
			for files in file_zip.namelist():
				file_zip.extract(files, os.path.join(self.path, zipdir))  # 解压到指定文件目录
			file_zip.close()


class readword():
    def __init__(self, dir):
        self.dir = dir

    def wordsToHtml(self):
        word = win32.Dispatch('Word.Application')
        filelistmergecell = glob.glob(self.dir + '\*.doc')
        # print (filelistmergecell)
        for wordfullName in filelistmergecell:
            doc = word.Documents.Open(wordfullName)
            htmlfullName = wordfullName[:-3] + 'html'
            txtfullName = wordfullName[:-3] + 'txt'

            print('正在处理图片----------' + htmlfullName)
            print('正在处理文字----------' + txtfullName)

            doc.SaveAs(htmlfullName, 10)
            doc.SaveAs(txtfullName, 5)

            os.remove(htmlfullName)
            print('正在删除html文件----------' + htmlfullName)
            doc.Close()

        filelist2 = glob.glob(self.dir + '\*.docx')
        # print (filelist2)
        for wordfullName in filelist2:
            doc = word.Documents.Open(wordfullName)
            htmlfullName = wordfullName[:-4] + 'html'
            txtfullName = wordfullName[:-4] + 'txt'

            print('正在处理图片----------' + htmlfullName)
            print('正在处理文字----------' + txtfullName)

            doc.SaveAs(htmlfullName, 10)
            doc.SaveAs(txtfullName, 5)

            os.remove(htmlfullName)
            print('正在删除html文件----------' + htmlfullName)
            doc.Close()
        word.Quit()

    def gettable(self):
        d = Document(self.dir)
        n = len(d.tables)
        df = pd.DataFrame([], columns=['参与门店', '活动时间', '活动要求', '活动支持'], index=range(0, n))
        for i in range(0, n):
            df['参与门店'][i] = d.tables[i].cell(0, 1).text
            df['活动时间'][i] = d.tables[i].cell(0, 3).text
            df['活动要求'][i] = d.tables[i].cell(0, 5).text
            df['活动支持'][i] = d.tables[i].cell(0, 7).text

        # df.to_csv(os.path.join(self.dir, f_name + '.csv'), encoding='gbk',index = False)
        return df


if __name__ == '__main__':
    # 读取MySQL中关于导入文件记录的文件
 #   filerecord = pd.read_sql(sqlcmd, conn)

    # 阿里云地址
    path = r'C:\Users\yunrui.hu\Desktop\t'
    filenames = os.listdir(path)
    for filename in filenames:
        # for i in range(len(filerecord[ ])):
        #     if str(filerecord[][i]).strip() == filename[0]:
        #         unit_name = filerecord[][i]
     # if os.path.splitext(filename)[1] in ('.doc', '.docx') and 'WTD' in str(
     #           os.path.splitext(filename)[0]) and unit_name == '武汉创洁工贸有限公司':
        if os.path.splitext(filename)[1] in ('.doc', '.docx') and 'WTD' in str(
            os.path.splitext(filename)[0]):
            rwo = readword(path)
            rwo.wordsToHtml()
            rwo.gettable()
        elif os.path.splitext(filename)[1] in ('.xls', '.xlsx') and '活动反馈' in str(
            os.path.splitext(filename)[0]):
            rex = readexcel(path)
            rex.read_marketingfile()
        elif os.path.splitext(filename)[1] in ('.xls', '.xlsx') and '一页纸' in str(
            os.path.splitext(filename)[0]):
            rex = readexcel(path)
            rex.re_xls_xlsx()
            rex.getexceldata('sheet1')
            rex.chage_file_name()
            rex.unzip_file()
        else:
            print('#############文件夹中有其他格式文件###############')